import json
import io
import tempfile

import pandas as pd

from utils.ocr import (
    extract_text_from_image
)

from langchain_core.documents import (
    Document
)

from docx import (
    Document as WordDocument
)

from langchain_community.document_loaders import (
    PyPDFLoader,
    CSVLoader,
    TextLoader
)


def load_document(
    uploaded_file
):

    file_name = uploaded_file.name.lower()

    suffix = "." + file_name.split(
        "."
    )[-1]

    # DOCX
    if file_name.endswith(
        ".docx"
    ):

        file_stream = io.BytesIO(
            uploaded_file.getvalue()
        )

        doc = WordDocument(
            file_stream
        )

        text = "\n".join(
            [
                para.text
                for para in doc.paragraphs
                if para.text.strip()
            ]
        )

        return [
            Document(
                page_content=text
            )
        ]

    # Remaining files
    with tempfile.NamedTemporaryFile(
        delete=False,
        suffix=suffix
    ) as tmp_file:

        tmp_file.write(
            uploaded_file.getvalue()
        )

        file_path = tmp_file.name

    # PDF
    if file_name.endswith(
        ".pdf"
    ):

        return PyPDFLoader(
            file_path
        ).load()

    # CSV
    elif file_name.endswith(
        ".csv"
    ):

        return CSVLoader(
            file_path
        ).load()

    # TXT
    elif file_name.endswith(
        ".txt"
    ):

        return TextLoader(
            file_path
        ).load()

    # JSON
    elif file_name.endswith(
        ".json"
    ):

        with open(
            file_path,
            "r",
            encoding="utf-8"
        ) as f:

            data = json.load(
                f
            )

        return [
            Document(
                page_content=json.dumps(
                    data,
                    indent=2
                )
            )
        ]

    # XLSX
    elif file_name.endswith(
        ".xlsx"
    ):

        df = pd.read_excel(
            file_path
        )

        text = df.to_string(
            index=False
        )

        return [
            Document(
                page_content=text
            )
        ]

    # PNG / JPG / JPEG
    elif (
        file_name.endswith(".png")
        or file_name.endswith(".jpg")
        or file_name.endswith(".jpeg")
    ):

        text = extract_text_from_image(
            file_path
        )

        return [
            Document(
                page_content=text
            )
        ]

    else:

        raise ValueError(
            f"Unsupported file type: {suffix}"
        )